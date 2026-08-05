from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────

ORANGE = RGBColor(0xE8, 0x65, 0x14)   # Brand orange
DARK   = RGBColor(0x1F, 0x29, 0x37)   # Near-black
GRAY   = RGBColor(0x6B, 0x72, 0x80)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xFF, 0xF7, 0xED) # Warm white

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color or (ORANGE if level == 1 else DARK)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(20)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_para(doc, text, bold=False, italic=False, size=10.5, color=None, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_code_block(doc, lines):
    """Monospace shaded block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, 'E86514')
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = 'FFF7ED' if idx % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E86514')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_page_margins(doc, top=1, bottom=1, left=1.2, right=1.2):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_cover(doc, title, subtitle, version='v1.0', date='Julio 2026'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('🍕  PIZZA YAJA')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ORANGE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(subtitle)
    run3.font.size = Pt(13)
    run3.font.color.rgb = GRAY
    run3.font.italic = True

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f'{version}  ·  {date}')
    run4.font.size = Pt(10)
    run4.font.color.rgb = GRAY

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
# MANUAL TÉCNICO
# ══════════════════════════════════════════════════════════════════

def build_manual_tecnico():
    doc = Document()
    set_page_margins(doc)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    add_cover(doc, 'Manual Técnico', 'Sistema ERP — Pizza Yaja')

    # ── 1. ARQUITECTURA GENERAL ──────────────────────────────────
    add_heading(doc, '1. Arquitectura General')
    add_divider(doc)
    add_para(doc,
        'Pizza Yaja es una aplicación web monolítica compuesta por dos capas: '
        'una aplicación POS (Punto de Venta) para el personal operativo y un módulo ERP '
        'para la administración del negocio. Ambas capas comparten la misma base de datos '
        'MySQL y el mismo servidor PHP.')

    add_heading(doc, '1.1 Diagrama de capas', level=2)
    add_code_block(doc, [
        '┌──────────────────────────────────────────────────────┐',
        '│              NAVEGADOR (Cliente)                     │',
        '│   HTML5 · CSS3 · JS · Chart.js 4.4.0               │',
        '└────────────────────┬─────────────────────────────────┘',
        '                     │ HTTP/AJAX',
        '┌────────────────────▼─────────────────────────────────┐',
        '│              SERVIDOR WEB — Apache (XAMPP)           │',
        '│   PHP 7+  ·  Sesiones PHP  ·  PDO                   │',
        '│                                                      │',
        '│   POS App           ERP Dashboard                    │',
        '│   /index.php        /erp/index.php                   │',
        '│   /menu.php         /erp/ventas.php                  │',
        '│   /cocina.php       /erp/inventario.php  …           │',
        '│   /pago.php                                          │',
        '└────────────────────┬─────────────────────────────────┘',
        '                     │ PDO / MySQL',
        '┌────────────────────▼─────────────────────────────────┐',
        '│              MySQL 8 (puerto 3307)                   │',
        '│              Base de datos: pizzeria                 │',
        '└──────────────────────────────────────────────────────┘',
    ])

    add_heading(doc, '1.2 Flujo de autenticación', level=2)
    add_para(doc, 'El login (/index.php) valida las credenciales contra la tabla usuarios. '
        'Con éxito se guarda en la sesión PHP: usuario_id, nombre y rol. '
        'Cada página protegida verifica la sesión y redirige según el rol:')
    add_bullet(doc, 'admin  →  /erp/index.php  (Dashboard ERP completo)')
    add_bullet(doc, 'camarero  →  /ordenes_activas.php  y  /menu.php')
    add_bullet(doc, 'cocina  →  /cocina.php  (pantalla de cocina)')

    # ── 2. STACK TECNOLÓGICO ─────────────────────────────────────
    add_heading(doc, '2. Stack Tecnológico')
    add_divider(doc)

    add_table(doc,
        ['Capa', 'Tecnología', 'Versión / Detalle'],
        [
            ['Backend',    'PHP',          '7+ (sin framework, PDO nativo)'],
            ['Servidor web','Apache',       'XAMPP (Windows local)'],
            ['Base de datos','MySQL / MariaDB','Puerto 3307, charset utf8mb4'],
            ['Frontend',   'HTML5 / CSS3', 'Vanilla — sin framework CSS'],
            ['JavaScript', 'Vanilla JS',   'ES6+, fetch API para AJAX'],
            ['Gráficas',   'Chart.js',     '4.4.0 (CDN)'],
            ['Íconos',     'Emoji / CSS',  'Sin librería de íconos externa'],
        ],
        col_widths=[1.5, 1.8, 3.0]
    )

    add_para(doc, 'No existe archivo package.json ni composer.json. '
        'Las únicas dependencias externas se cargan desde CDN (Chart.js).')

    # ── 3. ESTRUCTURA DE CARPETAS ────────────────────────────────
    add_heading(doc, '3. Estructura de Carpetas')
    add_divider(doc)

    add_code_block(doc, [
        'C:\\xampp\\htdocs\\Pizzeria\\',
        '├── config.php                   # Conexión a BD (host, puerto, credenciales)',
        '├── index.php                    # Login / autenticación',
        '├── admin.php                    # Dashboard admin del POS',
        '├── menu.php                     # Pantalla de pedidos (camarero)',
        '├── cocina.php                   # Vista de cocina (rol cocina)',
        '├── ordenes_activas.php          # Listado de órdenes activas',
        '├── pago.php                     # Procesamiento de pagos',
        '├── finalizar_orden.php          # Cierre de orden',
        '├── detalle_orden_ajax.php       # API AJAX: detalle de orden (JSON)',
        '├── actualizar_estado_item.php   # API AJAX: estado de ítem de orden',
        '├── agregar_item_orden.php       # API AJAX: agregar ítem a orden',
        '└── erp/                         # ◀ Módulo ERP (requiere rol admin)',
        '    ├── index.php                # Dashboard KPI',
        '    ├── ventas.php               # Reporte de ventas',
        '    ├── historial.php            # Historial de órdenes',
        '    ├── inventario.php           # Gestión de inventario',
        '    ├── gastos.php               # Registro de gastos',
        '    ├── productos.php            # Gestión de productos/menú',
        '    ├── usuarios.php             # Gestión de usuarios',
        '    ├── docs/                    # ◀ Documentación (este directorio)',
        '    └── includes/',
        '        └── sidebar.php          # Sidebar compartido + estilos globales ERP',
    ])

    # ── 4. ENTORNO LOCAL ─────────────────────────────────────────
    add_heading(doc, '4. Cómo levantar el entorno local')
    add_divider(doc)

    add_heading(doc, '4.1 Requisitos previos', level=2)
    add_bullet(doc, 'XAMPP instalado con Apache y MySQL activos')
    add_bullet(doc, 'PHP 7.4+ (incluido en XAMPP)')
    add_bullet(doc, 'MySQL corriendo en puerto 3307 (revisar config.php si difiere)')
    add_bullet(doc, 'Navegador moderno (Chrome / Firefox)')

    add_heading(doc, '4.2 Pasos de instalación', level=2)

    steps = [
        ('Clonar / copiar el proyecto',
         'Colocar la carpeta Pizzeria/ dentro de C:\\xampp\\htdocs\\'),
        ('Iniciar XAMPP',
         'Abrir XAMPP Control Panel → Start Apache y MySQL'),
        ('Crear la base de datos',
         'Acceder a phpMyAdmin (http://localhost/phpmyadmin) → Nueva BD → pizzeria → utf8mb4_unicode_ci'),
        ('Importar el esquema SQL',
         'Importar el dump pizzeria.sql si existe, o dejar que las tablas se creen automáticamente al primer uso.'),
        ('Verificar config.php',
         'Confirmar que host, puerto y nombre de BD coinciden con el entorno.'),
        ('Acceder a la app',
         'Abrir http://localhost/Pizzeria/ en el navegador.'),
        ('Crear usuario admin',
         'Si la BD está vacía, insertar manualmente un usuario con rol admin en la tabla usuarios.'),
    ]
    for i, (title, detail) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f'  {i}.  {title}: ')
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r2 = p.add_run(detail)
        r2.font.size = Pt(10.5)

    add_heading(doc, '4.3 Creación manual de usuario admin (SQL)', level=2)
    add_code_block(doc, [
        "INSERT INTO usuarios (nombre, usuario, contrasena, rol, activo)",
        "VALUES ('Administrador', 'admin', 'tu_password', 'admin', 1);",
        '',
        "-- ⚠️  Las contraseñas están en texto plano actualmente.",
        "-- Se recomienda implementar password_hash() en PHP antes de producción.",
    ])

    # ── 5. VARIABLES DE ENTORNO ──────────────────────────────────
    add_heading(doc, '5. Variables de Entorno / Configuración')
    add_divider(doc)
    add_para(doc,
        'El proyecto no usa un archivo .env. La configuración se centraliza en config.php. '
        'A continuación se listan las variables que deben ajustarse por entorno:')

    add_table(doc,
        ['Variable (config.php)', 'Descripción', 'Valor por defecto'],
        [
            ['$host',     'Host del servidor MySQL',             'localhost'],
            ['$port',     'Puerto MySQL',                        '3307'],
            ['$dbname',   'Nombre de la base de datos',         'pizzeria'],
            ['$username', 'Usuario de MySQL',                    'root'],
            ['$password', 'Contraseña de MySQL',                 '(vacía)'],
            ['$timezone', 'Zona horaria para fechas',            'America/Costa_Rica'],
        ],
        col_widths=[2.0, 2.5, 2.0]
    )
    add_para(doc,
        'Recomendación: migrar estos valores a un archivo .env '
        'y cargarlo con vlucas/phpdotenv antes de pasar a producción.',
        italic=True, color=GRAY)

    # ── 6. MODELO DE DATOS ───────────────────────────────────────
    add_heading(doc, '6. Modelo de Datos')
    add_divider(doc)
    add_para(doc, 'Base de datos: pizzeria | charset: utf8mb4 | colación: utf8mb4_unicode_ci')

    tables_info = [
        ('ordenes', 'Registro principal de cada orden/pedido', [
            ('id', 'INT PK AI', 'Identificador único'),
            ('numero_orden', 'VARCHAR', 'Número visible para el usuario'),
            ('nombre_cliente', 'VARCHAR', 'Nombre del cliente'),
            ('total', 'DECIMAL(10,2)', 'Total de la orden'),
            ('estado', 'ENUM', 'pendiente | en_proceso | listo | pagado'),
            ('fecha_creacion', 'DATETIME', 'Timestamp de creación'),
            ('tipo_servicio', 'VARCHAR', 'para_llevar | en_sitio'),
            ('metodo_pago', 'VARCHAR', 'efectivo | sinpe | tarjeta'),
        ]),
        ('detalle_orden', 'Líneas de ítem por orden', [
            ('id', 'INT PK AI', ''),
            ('orden_id', 'INT FK → ordenes.id', ''),
            ('producto_nombre', 'VARCHAR', 'Nombre al momento de la venta'),
            ('cantidad', 'INT', ''),
            ('precio_unitario', 'DECIMAL(10,2)', ''),
        ]),
        ('pagos', 'Registro de pagos (soporta pagos mixtos)', [
            ('id', 'INT PK AI', ''),
            ('orden_id', 'INT FK → ordenes.id', ''),
            ('metodo_pago', 'VARCHAR', 'efectivo | sinpe | tarjeta'),
            ('monto_aplicado', 'DECIMAL(10,2)', ''),
        ]),
        ('productos', 'Catálogo de productos del menú', [
            ('id', 'INT PK AI', ''),
            ('nombre', 'VARCHAR', ''),
            ('descripcion', 'TEXT', ''),
            ('precio', 'DECIMAL(10,2)', ''),
            ('categoria', 'VARCHAR', 'Texto legible de categoría'),
            ('disponible', 'TINYINT(1)', '1 = activo, 0 = inactivo'),
            ('categoria_id', 'INT FK → categorias.id', ''),
        ]),
        ('categorias', 'Categorías del menú', [
            ('id', 'INT PK AI', ''),
            ('nombre', 'VARCHAR', ''),
            ('orden', 'INT', 'Orden de aparición en el menú'),
        ]),
        ('usuarios', 'Cuentas de usuario del sistema', [
            ('id', 'INT PK AI', ''),
            ('nombre', 'VARCHAR', 'Nombre completo'),
            ('usuario', 'VARCHAR UNIQUE', 'Login'),
            ('contrasena', 'VARCHAR', '⚠️ Texto plano actualmente'),
            ('rol', 'ENUM', 'admin | camarero | cocina'),
            ('activo', 'TINYINT(1)', '1 = habilitado'),
        ]),
        ('gastos', 'Registro de gastos del negocio (creada automáticamente)', [
            ('id', 'INT PK AI', ''),
            ('concepto', 'VARCHAR', 'Descripción del gasto'),
            ('categoria', 'ENUM', 'ingredientes | servicios | personal | equipo | marketing | otro'),
            ('monto', 'DECIMAL(10,2)', ''),
            ('fecha', 'DATE', ''),
            ('nota', 'TEXT', 'Observaciones'),
            ('usuario_id', 'INT FK → usuarios.id', ''),
            ('created_at', 'DATETIME', ''),
        ]),
        ('ingredientes', 'Inventario de insumos (creada automáticamente)', [
            ('id', 'INT PK AI', ''),
            ('nombre', 'VARCHAR', ''),
            ('unidad', 'VARCHAR', 'kg, litros, unidades…'),
            ('stock_actual', 'DECIMAL(10,3)', ''),
            ('stock_minimo', 'DECIMAL(10,3)', 'Umbral de alerta'),
            ('costo_unitario', 'DECIMAL(10,2)', ''),
            ('activo', 'TINYINT(1)', ''),
            ('created_at', 'DATETIME', ''),
        ]),
        ('movimientos_inventario', 'Auditoría de movimientos de stock', [
            ('id', 'INT PK AI', ''),
            ('ingrediente_id', 'INT FK → ingredientes.id', ''),
            ('tipo', 'ENUM', 'entrada | salida | ajuste'),
            ('cantidad', 'DECIMAL(10,3)', ''),
            ('stock_antes', 'DECIMAL(10,3)', ''),
            ('stock_despues', 'DECIMAL(10,3)', ''),
            ('nota', 'TEXT', ''),
            ('usuario_id', 'INT FK → usuarios.id', ''),
            ('fecha', 'DATETIME', ''),
        ]),
    ]

    for tname, tdesc, cols in tables_info:
        add_heading(doc, f'Tabla: {tname}', level=2)
        add_para(doc, tdesc, italic=True, color=GRAY)
        add_table(doc,
            ['Columna', 'Tipo', 'Descripción'],
            [[c[0], c[1], c[2]] for c in cols],
            col_widths=[1.8, 2.0, 2.8]
        )

    # ── 7. ENDPOINTS ─────────────────────────────────────────────
    add_heading(doc, '7. Endpoints Principales')
    add_divider(doc)

    add_heading(doc, '7.1 POS — Aplicación Operativa', level=2)
    add_table(doc,
        ['Ruta', 'Método', 'Rol', 'Descripción'],
        [
            ['/Pizzeria/',                      'GET',      'todos',      'Login / selección de rol'],
            ['/menu.php',                       'GET/POST', 'camarero',   'Crear/editar órdenes'],
            ['/ordenes_activas.php',            'GET',      'camarero',   'Ver órdenes en curso'],
            ['/cocina.php',                     'GET',      'cocina',     'Panel de cocina'],
            ['/pago.php?orden_id=X',            'GET/POST', 'camarero',   'Procesar pago de orden'],
            ['/finalizar_orden.php',            'POST',     'camarero',   'Cerrar orden como pagada'],
            ['/detalle_orden_ajax.php',         'GET',      'camarero',   'JSON: detalle de orden'],
            ['/actualizar_estado_item.php',     'POST',     'cocina',     'JSON: cambiar estado ítem'],
            ['/agregar_item_orden.php',         'POST',     'camarero',   'JSON: añadir ítem'],
        ],
        col_widths=[2.4, 1.0, 1.2, 2.0]
    )

    add_heading(doc, '7.2 ERP — Panel de Administración (requiere rol admin)', level=2)
    add_table(doc,
        ['Ruta', 'Params GET', 'Acciones AJAX (POST)', 'Descripción'],
        [
            ['/erp/index.php',       '—',                '—',                              'Dashboard KPI'],
            ['/erp/ventas.php',      'dias=7|30|90',     '—',                              'Reporte de ventas'],
            ['/erp/historial.php',   'q, fecha, p',      '—',                              'Historial de órdenes'],
            ['/erp/inventario.php',  'tab=stock|movs',   'crear, editar, movimiento, eliminar', 'Inventario'],
            ['/erp/gastos.php',      'dias=7|30|90',     'crear, editar, eliminar',        'Gastos'],
            ['/erp/productos.php',   '—',                'crear, editar, toggle, eliminar','Productos/Menú'],
            ['/erp/usuarios.php',    '—',                'crear, editar, eliminar',        'Usuarios'],
        ],
        col_widths=[1.8, 1.4, 2.2, 1.8]
    )

    add_para(doc,
        'Todas las acciones AJAX responden con JSON { success: bool, message: str }. '
        'Los formularios se envían vía fetch() desde el frontend.',
        italic=True, color=GRAY)

    # ── 8. SEGURIDAD ─────────────────────────────────────────────
    add_heading(doc, '8. Notas de Seguridad y Deuda Técnica')
    add_divider(doc)

    add_para(doc, 'Puntos críticos a resolver antes de producción:', bold=True)
    items = [
        ('Contraseñas en texto plano',
         'La tabla usuarios almacena contraseñas sin hashear. Implementar password_hash() / password_verify() en PHP.'),
        ('Sin archivo .env',
         'Las credenciales de BD están en config.php. Mover a variables de entorno.'),
        ('Sin CSRF protection',
         'Los formularios AJAX no usan tokens CSRF. Añadir token en sesión y verificarlo en cada POST.'),
        ('Sin prepared statements en algunos archivos',
         'Algunos archivos del POS usan interpolación directa de variables en queries. Migrar a PDO bindParam().'),
        ('Sin HTTPS',
         'El proyecto corre sobre HTTP local. Para producción, configurar SSL/TLS.'),
        ('Tablas auto-creadas en producción',
         'gastos.php e inventario.php crean sus tablas con CREATE TABLE IF NOT EXISTS en cada request. Migrar a un script de migración.'),
    ]
    for title, detail in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f'  ⚠  {title}: ')
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        r2 = p.add_run(detail)
        r2.font.size = Pt(10.5)

    path = os.path.join(OUTPUT_DIR, 'MANUAL_TECNICO.docx')
    doc.save(path)
    print(f'✅  MANUAL_TECNICO.docx guardado en {path}')


# ══════════════════════════════════════════════════════════════════
# MANUAL DE USUARIO
# ══════════════════════════════════════════════════════════════════

def build_manual_usuario():
    doc = Document()
    set_page_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    add_cover(doc, 'Manual de Usuario', 'Sistema ERP — Pizza Yaja')

    # ── INTRO ────────────────────────────────────────────────────
    add_heading(doc, 'Bienvenido al Sistema Pizza Yaja')
    add_divider(doc)
    add_para(doc,
        'Este manual explica cómo usar el sistema de gestión de Pizza Yaja. '
        'El sistema tiene dos partes: la aplicación de pedidos (para cajeros y cocina) '
        'y el panel ERP (para el administrador). '
        'Siga las instrucciones según su rol en el restaurante.')

    add_heading(doc, 'Roles del Sistema', level=2)
    add_table(doc,
        ['Rol', 'Quién lo usa', 'Qué puede hacer'],
        [
            ['Administrador', 'Dueño / Gerente', 'Ver reportes, gestionar productos, inventario, gastos y usuarios'],
            ['Camarero',      'Mesero / Cajero', 'Tomar pedidos, ver órdenes activas, cobrar'],
            ['Cocina',        'Cocinero',        'Ver los pedidos entrantes y marcarlos como listos'],
        ],
        col_widths=[1.5, 2.0, 3.2]
    )

    # ── ACCESO ───────────────────────────────────────────────────
    add_heading(doc, '1. Cómo Ingresar al Sistema')
    add_divider(doc)

    add_para(doc, 'Paso 1: Abrir el navegador e ir a la dirección del sistema (su administrador le indicará la URL).')
    add_para(doc, 'Paso 2: Escribir su usuario y contraseña.')
    add_para(doc, 'Paso 3: Hacer clic en Ingresar.')
    add_para(doc,
        'El sistema lo llevará automáticamente a la pantalla correcta según su rol. '
        'No necesita recordar ninguna ruta especial.')

    add_para(doc, 'Si olvidó su contraseña, comuníquese con el administrador del sistema.', italic=True, color=GRAY)

    # ── SECCIÓN CAMARERO ─────────────────────────────────────────
    add_heading(doc, '2. Sección para Camareros / Cajeros')
    add_divider(doc)

    add_heading(doc, '2.1 Tomar un Pedido Nuevo', level=2)
    add_para(doc,
        'Desde el menú principal, verá todos los productos organizados por categoría '
        '(Pizzas, Bebidas, Entradas, etc.).')

    steps_camarero = [
        'Busque el producto que el cliente quiere y haga clic en él.',
        'Se agregará al carrito de la derecha.',
        'Puede cambiar la cantidad usando los botones + y −.',
        'Cuando el pedido esté completo, haga clic en Enviar Orden.',
        'El pedido aparecerá automáticamente en la pantalla de cocina.',
    ]
    for i, s in enumerate(steps_camarero, 1):
        add_bullet(doc, f'{i}. {s}')

    add_heading(doc, '2.2 Ver Órdenes Activas', level=2)
    add_para(doc,
        'En la pantalla de Órdenes Activas verá todas las órdenes en curso con su estado actual:')
    add_bullet(doc, 'Pendiente — recién enviada a cocina')
    add_bullet(doc, 'En proceso — cocina está preparando el pedido')
    add_bullet(doc, 'Listo — el pedido está listo para entregar o cobrar')
    add_para(doc, 'La pantalla se actualiza sola, no necesita recargar la página.')

    add_heading(doc, '2.3 Cobrar un Pedido', level=2)
    steps_pago = [
        'En la lista de órdenes activas, haga clic en el botón Cobrar de la orden correspondiente.',
        'Seleccione el método de pago: Efectivo, SINPE Móvil o Tarjeta.',
        'Si el cliente paga con varios métodos (por ejemplo, parte en efectivo y parte con SINPE), puede registrar ambos pagos.',
        'Haga clic en Finalizar Pago.',
        'La orden quedará marcada como pagada y desaparecerá de la lista activa.',
    ]
    for i, s in enumerate(steps_pago, 1):
        add_bullet(doc, f'{i}. {s}')

    add_para(doc,
        'Consejo: Si el cliente paga en efectivo, el sistema calcula automáticamente el vuelto.',
        italic=True, color=GRAY)

    # ── SECCIÓN COCINA ───────────────────────────────────────────
    add_heading(doc, '3. Sección para el Personal de Cocina')
    add_divider(doc)

    add_heading(doc, '3.1 Pantalla de Cocina', level=2)
    add_para(doc,
        'Al ingresar, verá la pantalla de cocina con los pedidos entrantes en tiempo real. '
        'Cada tarjeta muestra el número de orden, el nombre del cliente y los ítems a preparar.')

    add_heading(doc, '3.2 Marcar Ítems como Listos', level=2)
    steps_cocina = [
        'Cuando termine de preparar un ítem, haga clic en el botón Listo junto a ese ítem.',
        'El ítem cambiará de color para indicar que está listo.',
        'Cuando todos los ítems de una orden estén listos, la orden completa se marca como Lista.',
        'El camarero verá el cambio en su pantalla automáticamente.',
    ]
    for i, s in enumerate(steps_cocina, 1):
        add_bullet(doc, f'{i}. {s}')

    add_para(doc, 'No necesita usar ningún otro botón ni navegar a otras páginas.', italic=True, color=GRAY)

    # ── SECCIÓN ADMIN ────────────────────────────────────────────
    add_heading(doc, '4. Sección para el Administrador (Panel ERP)')
    add_divider(doc)
    add_para(doc,
        'Al ingresar con el rol de administrador, accede al panel ERP con el menú lateral izquierdo. '
        'Desde ahí puede navegar por todos los módulos de gestión.')

    add_heading(doc, '4.1 Dashboard (Inicio)', level=2)
    add_para(doc,
        'La pantalla de inicio muestra un resumen del negocio del día y la semana:')
    add_bullet(doc, 'Ventas del día: total de dinero cobrado hoy')
    add_bullet(doc, 'Número de órdenes: cantidad de pedidos atendidos')
    add_bullet(doc, 'Ticket promedio: gasto promedio por cliente')
    add_bullet(doc, 'Gráfica de ventas de los últimos 7 días')
    add_bullet(doc, 'Desglose por método de pago (efectivo, SINPE, tarjeta)')
    add_bullet(doc, 'Top 5 productos más vendidos')
    add_bullet(doc, 'Últimas órdenes del día')
    add_para(doc, 'Esta pantalla se actualiza cada vez que la recarga.', italic=True, color=GRAY)

    add_heading(doc, '4.2 Módulo de Ventas', level=2)
    add_para(doc,
        'Aquí puede ver análisis detallados de las ventas con filtros de tiempo:')
    add_bullet(doc, 'Últimos 7 días, 30 días o 90 días')
    add_bullet(doc, 'Ventas por día de la semana (para saber los días más ocupados)')
    add_bullet(doc, 'Ventas por método de pago')
    add_bullet(doc, 'Horas pico del día (para planificar el personal)')
    add_bullet(doc, 'Top 10 productos más vendidos en el período')
    add_para(doc,
        'Use este módulo para decidir qué productos promocionar, cuándo necesita más personal '
        'y cuáles son los mejores días de la semana.', italic=True, color=GRAY)

    add_heading(doc, '4.3 Historial de Órdenes', level=2)
    add_para(doc, 'Consulte cualquier orden pasada. Puede buscar por:')
    add_bullet(doc, 'Número de orden')
    add_bullet(doc, 'Nombre del cliente')
    add_bullet(doc, 'Fecha específica')
    add_para(doc,
        'Haga clic en el botón Ver detalle de cualquier orden para ver todos los productos '
        'que se pidieron, el método de pago y la hora exacta.')

    add_heading(doc, '4.4 Inventario', level=2)
    add_para(doc,
        'Lleve control de los ingredientes e insumos del restaurante.')

    add_heading(doc, 'Ver el stock actual:', level=3)
    add_para(doc,
        'En la pestaña Stock verá todos los ingredientes con su cantidad actual. '
        'Los ingredientes en rojo o con alerta están por debajo del mínimo y necesitan reabastecerse.')

    add_heading(doc, 'Registrar un movimiento de inventario:', level=3)
    steps_inv = [
        'Haga clic en el botón Movimiento junto al ingrediente.',
        'Seleccione el tipo: Entrada (llegó mercadería), Salida (se usó) o Ajuste (corrección manual).',
        'Ingrese la cantidad y una nota opcional.',
        'Haga clic en Guardar.',
    ]
    for i, s in enumerate(steps_inv, 1):
        add_bullet(doc, f'{i}. {s}')

    add_heading(doc, 'Agregar un ingrediente nuevo:', level=3)
    add_para(doc,
        'Haga clic en Nuevo Ingrediente, complete el formulario con nombre, '
        'unidad de medida (kg, litros, unidades), stock inicial, stock mínimo y costo unitario.')

    add_heading(doc, '4.5 Gastos', level=2)
    add_para(doc,
        'Registre todos los gastos del negocio para calcular la ganancia real.')

    add_heading(doc, 'Registrar un gasto:', level=3)
    steps_gasto = [
        'Haga clic en Nuevo Gasto.',
        'Escriba el concepto (descripción del gasto).',
        'Seleccione la categoría: Ingredientes, Servicios (luz/agua/internet), Personal, Equipo, Marketing u Otro.',
        'Ingrese el monto y la fecha.',
        'Agregue una nota si lo desea.',
        'Haga clic en Guardar.',
    ]
    for i, s in enumerate(steps_gasto, 1):
        add_bullet(doc, f'{i}. {s}')

    add_para(doc,
        'El sistema calcula automáticamente el margen de ganancia comparando '
        'las ventas del período con los gastos registrados.', italic=True, color=GRAY)

    add_heading(doc, '4.6 Gestión de Productos (Menú)', level=2)
    add_para(doc,
        'Administre los productos que aparecen en el menú de la aplicación de pedidos.')

    add_heading(doc, 'Agregar un producto:', level=3)
    steps_prod = [
        'Haga clic en Nuevo Producto.',
        'Ingrese el nombre, descripción, precio y categoría.',
        'Haga clic en Guardar.',
        'El producto aparecerá inmediatamente en el menú de pedidos.',
    ]
    for i, s in enumerate(steps_prod, 1):
        add_bullet(doc, f'{i}. {s}')

    add_heading(doc, 'Activar o desactivar un producto:', level=3)
    add_para(doc,
        'Use el botón de disponibilidad (activo/inactivo) para ocultar temporalmente '
        'un producto del menú sin eliminarlo. Útil cuando se acaba un ingrediente.')

    add_heading(doc, 'Editar o eliminar:', level=3)
    add_para(doc,
        'Use los botones Editar y Eliminar en cada fila. '
        'Tenga cuidado al eliminar — la acción no se puede deshacer.')

    add_heading(doc, '4.7 Gestión de Usuarios', level=2)
    add_para(doc,
        'Cree y administre las cuentas del personal que usa el sistema.')

    add_heading(doc, 'Crear un usuario nuevo:', level=3)
    steps_usr = [
        'Haga clic en Nuevo Usuario.',
        'Ingrese el nombre completo, nombre de usuario (login) y contraseña.',
        'Seleccione el rol: Administrador, Camarero o Cocina.',
        'Haga clic en Guardar.',
    ]
    for i, s in enumerate(steps_usr, 1):
        add_bullet(doc, f'{i}. {s}')

    add_heading(doc, 'Desactivar un usuario:', level=3)
    add_para(doc,
        'Si un empleado deja de trabajar, use el botón Editar y cambie su estado a Inactivo. '
        'Así el usuario no podrá ingresar sin que sus datos se pierdan.')

    # ── FLUJOS TÍPICOS ───────────────────────────────────────────
    add_heading(doc, '5. Flujos Típicos del Día a Día')
    add_divider(doc)

    flows = [
        ('Apertura del turno',
         ['El administrador revisa el Dashboard para ver el resumen del día anterior.',
          'Verifica el inventario para detectar ingredientes por reponer.',
          'El personal de cocina y camareros ingresan con sus cuentas respectivas.']),
        ('Atención de un cliente (flujo completo)',
         ['Camarero toma el pedido desde /menu.php y lo envía.',
          'Cocina ve el pedido en /cocina.php y lo prepara.',
          'Cocina marca cada ítem como listo.',
          'Camarero entrega el pedido y procede al cobro.',
          'Camarero selecciona método de pago y finaliza la orden.',
          'La orden desaparece de activas y queda en el historial.']),
        ('Registro de gastos semanal',
         ['Administrador va al módulo Gastos.',
          'Registra facturas de proveedores (categoría: Ingredientes).',
          'Registra servicios (luz, agua, internet).',
          'Revisa el margen de ganancia calculado automáticamente.']),
        ('Actualización del menú',
         ['Administrador va al módulo Productos.',
          'Edita precios según necesidad.',
          'Desactiva temporalmente productos no disponibles.',
          'Los cambios se reflejan de inmediato para los camareros.']),
    ]

    for flow_title, flow_steps in flows:
        add_heading(doc, flow_title, level=2)
        for s in flow_steps:
            add_bullet(doc, s)

    # ── PREGUNTAS FRECUENTES ─────────────────────────────────────
    add_heading(doc, '6. Preguntas Frecuentes')
    add_divider(doc)

    faqs = [
        ('¿Qué hago si una orden fue cobrada por error?',
         'Comuníquese con el administrador. Actualmente el sistema no tiene función de reverso de pago; '
         'se debe registrar manualmente como un gasto en el módulo de Gastos.'),
        ('¿Puedo cambiar el precio de un producto en medio de un turno?',
         'Sí. Los cambios en Gestión de Productos se aplican inmediatamente. '
         'Sin embargo, los pedidos ya enviados no se recalculan.'),
        ('¿Qué pasa si la pantalla de cocina no muestra los pedidos?',
         'Recargar la página suele resolver el problema. Si persiste, verificar la conexión de red.'),
        ('¿Cómo sé cuáles ingredientes necesito comprar?',
         'En el módulo Inventario, los ingredientes con alerta de stock mínimo aparecen destacados. '
         'Esos son los que debe reponer.'),
        ('¿Puedo generar un reporte en PDF o Excel?',
         'Actualmente el sistema no tiene exportación directa. '
         'Puede usar la función de impresión del navegador (Ctrl+P) para guardar como PDF.'),
    ]

    for q, a in faqs:
        add_heading(doc, q, level=2)
        add_para(doc, a)

    path = os.path.join(OUTPUT_DIR, 'MANUAL_USUARIO.docx')
    doc.save(path)
    print(f'✅  MANUAL_USUARIO.docx guardado en {path}')


if __name__ == '__main__':
    build_manual_tecnico()
    build_manual_usuario()
    print('\n✅  Ambos documentos generados correctamente.')
